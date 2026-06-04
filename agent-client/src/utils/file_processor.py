"""
文件处理器模块

支持PDF、Word、Excel、Markdown、TXT和代码文件的文本提取
"""

import os
import logging
from typing import Optional, List, Dict, Any
from abc import ABC, abstractmethod
from pathlib import Path

from utils.helpers import truncate_file_content

# 配置日志
logger = logging.getLogger(__name__)


class BaseFileProcessor(ABC):
    """文件处理器基类"""
    
    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """检查是否能处理该文件"""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """提取文本内容"""
        pass


class PDFProcessor(BaseFileProcessor):
    """PDF文件处理器"""
    
    SUPPORTED_EXTENSIONS = ['.pdf']
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        提取PDF文本内容
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            import PyPDF2
            
            text_parts = []
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(f"--- 第 {i + 1} 页 ---")
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"提取第 {i + 1} 页失败: {str(e)}")
                        text_parts.append(f"--- 第 {i + 1} 页 (提取失败) ---")
                
                # 尝试提取元数据
                try:
                    metadata = reader.metadata
                    if metadata:
                        meta_info = []
                        if metadata.get("/Title"):
                            meta_info.append(f"标题: {metadata['/Title']}")
                        if metadata.get("/Author"):
                            meta_info.append(f"作者: {metadata['/Author']}")
                        if metadata.get("/Creator"):
                            meta_info.append(f"创建者: {metadata['/Creator']}")
                        
                        if meta_info:
                            text_parts.insert(0, "=== 文档信息 ===")
                            text_parts.extend(meta_info)
                            text_parts.insert(len(meta_info) + 1, "")
                except Exception:
                    pass
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("PyPDF2未安装，请运行: pip install PyPDF2")
            return f"错误: 无法处理PDF文件，请安装PyPDF2库"
        except Exception as e:
            logger.error(f"PDF处理失败: {str(e)}")
            return f"错误: PDF处理失败 - {str(e)}"


class WordProcessor(BaseFileProcessor):
    """Word文档处理器"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        提取Word文档文本内容
        
        Args:
            file_path: Word文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            from docx import Document
            
            text_parts = []
            
            doc = Document(file_path)
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                text_parts.append("\n[表格]")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("[/表格]\n")
            
            # 提取页眉页脚
            try:
                for section in doc.sections:
                    if section.header.paragraphs:
                        header_text = " ".join([p.text for p in section.header.paragraphs if p.text])
                        if header_text:
                            text_parts.insert(0, f"页眉: {header_text}\n")
                    
                    if section.footer.paragraphs:
                        footer_text = " ".join([p.text for p in section.footer.paragraphs if p.text])
                        if footer_text:
                            text_parts.append(f"\n页脚: {footer_text}")
            except Exception:
                pass
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx未安装，请运行: pip install python-docx")
            return f"错误: 无法处理Word文件，请安装python-docx库"
        except Exception as e:
            logger.error(f"Word处理失败: {str(e)}")
            return f"错误: Word处理失败 - {str(e)}"


class ExcelProcessor(BaseFileProcessor):
    """Excel文件处理器"""
    
    SUPPORTED_EXTENSIONS = ['.xlsx', '.xls', '.csv']
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        提取Excel文本内容
        
        Args:
            file_path: Excel文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            import openpyxl
            
            text_parts = []
            
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            
            for sheet_name in wb.sheetnames:
                text_parts.append(f"\n=== 工作表: {sheet_name} ===")
                
                sheet = wb[sheet_name]
                
                # 读取前100行
                row_count = 0
                max_rows = 100
                
                for row in sheet.iter_rows(max_row=max_rows, values_only=True):
                    if row_count >= max_rows:
                        break
                    
                    # 过滤空行
                    if any(cell is not None and str(cell).strip() for cell in row):
                        row_text = " | ".join([
                            str(cell) if cell is not None else ""
                            for cell in row
                        ])
                        text_parts.append(row_text)
                        row_count += 1
                
                if sheet.max_row > max_rows:
                    text_parts.append(f"... (还有 {sheet.max_row - max_rows} 行)")
            
            wb.close()
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("openpyxl未安装，请运行: pip install openpyxl")
            return f"错误: 无法处理Excel文件，请安装openpyxl库"
        except Exception as e:
            logger.error(f"Excel处理失败: {str(e)}")
            return f"错误: Excel处理失败 - {str(e)}"


class TextProcessor(BaseFileProcessor):
    """文本文件处理器（Markdown、TXT等）"""
    
    SUPPORTED_EXTENSIONS = [
        '.txt', '.md', '.markdown', '.rtf',
        '.json', '.xml', '.yaml', '.yml',
        '.log', '.ini', '.cfg', '.conf'
    ]
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        提取文本文件内容
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            str: 文件内容
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    # 对于JSON文件，格式化输出
                    if Path(file_path).suffix.lower() == '.json':
                        import json
                        try:
                            data = json.loads(content)
                            content = json.dumps(data, indent=2, ensure_ascii=False)
                        except json.JSONDecodeError:
                            pass
                    
                    return content
                    
                except UnicodeDecodeError:
                    continue
            
            return f"错误: 无法解码文件，请检查文件编码"
            
        except Exception as e:
            logger.error(f"文本文件处理失败: {str(e)}")
            return f"错误: 文本文件处理失败 - {str(e)}"


class CodeProcessor(BaseFileProcessor):
    """代码文件处理器"""
    
    SUPPORTED_EXTENSIONS = [
        '.py', '.java', '.c', '.cpp', '.h', '.cs',
        '.js', '.ts', '.jsx', '.tsx', '.vue',
        '.html', '.css', '.scss', '.sass', '.less',
        '.go', '.rs', '.rb', '.php', '.swift', '.kt',
        '.sql', '.sh', '.bash', '.bat', '.ps1',
        '.r', '.m', '.scala', '.pl', '.lua',
        '.toml', '.gradle'
    ]
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        提取代码文件内容
        
        Args:
            file_path: 代码文件路径
            
        Returns:
            str: 代码内容（带文件信息）
        """
        try:
            # 获取文件信息
            file_name = Path(file_path).name
            file_ext = Path(file_path).suffix.lstrip('.')
            file_size = os.path.getsize(file_path)
            
            # 读取代码
            encodings = ['utf-8', 'gbk', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return f"错误: 无法读取代码文件"
            
            # 构建输出
            header = f"=== 代码文件: {file_name} ==="
            info = f"语言: {file_ext} | 大小: {file_size} 字节 | 行数: {content.count(chr(10)) + 1}"
            
            # 限制代码长度
            max_chars = 50000
            if len(content) > max_chars:
                content = content[:max_chars] + f"\n\n... [代码已截断，共 {len(content)} 字符]"
            
            return f"{header}\n{info}\n\n```{file_ext}\n{content}\n```"
            
        except Exception as e:
            logger.error(f"代码文件处理失败: {str(e)}")
            return f"错误: 代码文件处理失败 - {str(e)}"


class FileProcessor:
    """
    文件处理器
    
    统一的文件处理接口，自动选择合适的处理器
    """
    
    def __init__(self):
        """初始化文件处理器"""
        self.processors: List[BaseFileProcessor] = [
            PDFProcessor(),
            WordProcessor(),
            ExcelProcessor(),
            CodeProcessor(),
            TextProcessor(),
        ]
        
        logger.info("文件处理器初始化完成，支持: PDF, Word, Excel, 代码, 文本")
    
    def get_processor(self, file_path: str) -> Optional[BaseFileProcessor]:
        """
        获取合适的文件处理器
        
        Args:
            file_path: 文件路径
            
        Returns:
            BaseFileProcessor: 处理器实例，如果没有找到返回None
        """
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def extract_text(self, file_path: str, max_length: int = 50000) -> str:
        """
        提取文件文本内容
        
        Args:
            file_path: 文件路径
            max_length: 最大返回长度
            
        Returns:
            str: 提取的文本内容
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return f"错误: 文件不存在 - {file_path}"
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        max_size = 50 * 1024 * 1024  # 50MB
        
        if file_size > max_size:
            return f"错误: 文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"
        
        # 获取处理器
        processor = self.get_processor(file_path)
        
        if processor is None:
            return f"错误: 不支持的文件类型 - {Path(file_path).suffix}"
        
        # 提取文本
        try:
            text = processor.extract_text(file_path)
            
            # 截断处理
            if len(text) > max_length:
                text = truncate_file_content(text, max_length=max_length // 100, max_chars=max_length)
            
            return text
            
        except Exception as e:
            logger.error(f"文件处理失败 [{file_path}]: {str(e)}")
            return f"错误: 文件处理失败 - {str(e)}"
    
    def extract_text_with_info(self, file_path: str) -> Dict[str, Any]:
        """
        提取文件文本内容（带详细信息）
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文件信息和内容的字典
        """
        file_path = os.path.abspath(file_path)
        
        result = {
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "file_extension": Path(file_path).suffix.lstrip('.'),
            "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            "success": False,
            "content": "",
            "error": None,
            "processor_type": None
        }
        
        # 检查文件
        if not os.path.exists(file_path):
            result["error"] = "文件不存在"
            return result
        
        # 获取处理器
        processor = self.get_processor(file_path)
        
        if processor is None:
            result["error"] = f"不支持的文件类型: {result['file_extension']}"
            return result
        
        # 提取文本
        try:
            result["content"] = processor.extract_text(file_path)
            result["success"] = True
            result["processor_type"] = processor.__class__.__name__
            
        except Exception as e:
            result["error"] = str(e)
            logger.error(f"文件处理失败: {str(e)}")
        
        return result
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名列表"""
        extensions = set()
        
        for processor in self.processors:
            if hasattr(processor, 'SUPPORTED_EXTENSIONS'):
                extensions.update(processor.SUPPORTED_EXTENSIONS)
        
        return sorted(list(extensions))
    
    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持"""
        return self.get_processor(file_path) is not None


# 便捷函数
def get_file_processor() -> FileProcessor:
    """获取文件处理器实例"""
    return FileProcessor()


def extract_file_text(file_path: str) -> str:
    """快速提取文件文本"""
    processor = get_file_processor()
    return processor.extract_text(file_path)


# 测试代码
if __name__ == "__main__":
    print("=" * 50)
    print("文件处理器测试")
    print("=" * 50)
    
    # 创建处理器
    processor = FileProcessor()
    
    # 显示支持的文件类型
    print("\n支持的文件扩展名:")
    extensions = processor.get_supported_extensions()
    print(", ".join(extensions))
    
    # 测试文件路径
    test_files = [
        "test.pdf",
        "test.docx",
        "test.xlsx",
        "test.md",
        "test.txt",
        "test.py",
    ]
    
    print("\n文件类型检测:")
    for test_file in test_files:
        supported = processor.is_supported(test_file)
        status = "✓ 支持" if supported else "✗ 不支持"
        print(f"  {test_file}: {status}")
    
    # 测试不支持的文件类型
    unsupported = ["test.jpg", "test.png", "test.exe", "test.mp3"]
    print("\n不支持的文件类型:")
    for test_file in unsupported:
        supported = processor.is_supported(test_file)
        status = "✗ 不支持" if not supported else "✓ 支持"
        print(f"  {test_file}: {status}")
    
    print("\n" + "=" * 50)
    print("文件处理器初始化完成！")
    print("=" * 50)
